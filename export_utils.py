import re
from docx import Document
from docx.shared import Pt, RGBColor
from fpdf import FPDF
import io

def clean_markdown(text):
    """
    Cleans up markdown text for basic processing.
    """
    # 1. Remove wrapping code blocks if the entire text is wrapped or mixed
    # This regex looks for ```markdown ... ``` and keeps only the inner content
    match = re.search(r'```markdown\s*(.*?)\s*```', text, flags=re.DOTALL | re.IGNORECASE)
    if match:
        text = match.group(1)
    else:
        # Fallback: just remove standalone backticks lines if they exist
        text = text.replace("```markdown", "").replace("```", "")
        
    return text.strip()

def add_formatted_text(paragraph, text):
    """
    Parses text for bold markers (**text**) and adds runs to the paragraph.
    """
    # Split by bold markers
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold text
            content = part[2:-2]
            if content: # Avoid empty bold markers
                run = paragraph.add_run(content)
                run.bold = True
        else:
            # Normal text
            if part: # Avoid empty strings
                paragraph.add_run(part)

def markdown_to_docx(markdown_text):
    """
    Converts simple Markdown text to a DOCX file stream.
    """
    # Clean the input text first
    markdown_text = clean_markdown(markdown_text)
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    lines = markdown_text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            # Heading 1
            h = doc.add_heading('', level=1)
            add_formatted_text(h, line[2:])
            h.style.font.color.rgb = RGBColor(0, 0, 0) # Black
        elif line.startswith('## '):
            # Heading 2
            h = doc.add_heading('', level=2)
            add_formatted_text(h, line[3:])
            h.style.font.color.rgb = RGBColor(0, 0, 0)
        elif line.startswith('### '):
            # Heading 3
            h = doc.add_heading('', level=3)
            add_formatted_text(h, line[4:])
            h.style.font.color.rgb = RGBColor(0, 0, 0)
        elif line.startswith('- ') or line.startswith('* '):
            # Bullet point
            p = doc.add_paragraph('', style='List Bullet')
            add_formatted_text(p, line[2:])
        else:
            # Normal text
            p = doc.add_paragraph()
            add_formatted_text(p, line)

    # Save to memory stream
    docx_stream = io.BytesIO()
    doc.save(docx_stream)
    docx_stream.seek(0)
    return docx_stream

class PDF(FPDF):
    def header(self):
        pass

    def footer(self):
        self.set_y(-15)
        self.set_font('Helvetica', 'I', 8)
        self.cell(0, 10, f'Page {self.page_no()}', align='C')

def sanitize_for_latin1(text):
    """
    Sanitizes text to ensure it's compatible with Latin-1 encoding.
    Replaces common problematic characters with ASCII equivalents.
    """
    replacements = {
        '\u2013': '-',   # En dash
        '\u2014': '--',  # Em dash
        '\u2018': "'",   # Left single quote
        '\u2019': "'",   # Right single quote
        '\u201c': '"',   # Left double quote
        '\u201d': '"',   # Right double quote
        '\u2022': '-',   # Bullet
        '\u2026': '...', # Ellipsis
        '\u00a0': ' ',   # Non-breaking space
        '•': '-',        # Bullet
        '–': '-',        # En dash
        '—': '--',       # Em dash
    }
    
    for char, replacement in replacements.items():
        text = text.replace(char, replacement)
        
    # Replace any other characters that might cause issues
    return text.encode('latin-1', 'replace').decode('latin-1')

def write_text_with_links(pdf, text):
    """
    Writes text to PDF, detecting Markdown links [text](url) and creating clickable links.
    Handles line wrapping safely.
    """
    # Pattern to find [text](url)
    link_pattern = re.compile(r'\[(.*?)\]\((.*?)\)')
    
    # Process the text part by part
    matches = list(link_pattern.finditer(text))
    if not matches:
        pdf.multi_cell(0, 5, text)
        return

    # If there are links, we need to handle them carefully.
    # Since multi_cell with mixed styles is hard in FPDF, we'll strip links for simplicity
    # or just print the text.
    # A better approach for bullets with links is to just print the text representation
    # and maybe put the link in a footnote or just make the whole text clickable?
    # For now, let's just clean the link syntax to text for bullets to avoid layout issues,
    # unless it's a standalone link line.
    
    # Heuristic: If the line is short, try to render it with links.
    # If it's long (likely wrapping), stripping might be safer for alignment.
    
    # Actually, let's try to just render the text part of the link.
    clean_text = ""
    last_idx = 0
    for match in matches:
        clean_text += text[last_idx:match.start()]
        clean_text += match.group(1) # The link text
        last_idx = match.end()
    clean_text += text[last_idx:]
    
    pdf.multi_cell(0, 5, clean_text)

def draw_split_line(pdf, left_text, right_text, is_bold_left=False, is_italic_left=False):
    """
    Draws a line with text aligned to the left and text aligned to the right.
    Used for: Company Name ...... Date
              Role .......... Location
    """
    # Save current position
    y = pdf.get_y()
    
    # Set font for Left text
    style = ''
    if is_bold_left: style += 'B'
    if is_italic_left: style += 'I'
    pdf.set_font("Helvetica", style, 11) # Slightly larger for headers
    
    # Calculate width of right text
    right_width = pdf.get_string_width(right_text) + 2
    
    # 1. Write Right Text aligned to the right margin
    # We need to set x to page_width - right_margin - text_width
    page_width = pdf.w
    right_margin = pdf.r_margin
    x_right = page_width - right_margin - right_width
    
    pdf.set_xy(x_right, y)
    pdf.set_font("Helvetica", '', 11) # Date/Location usually normal font
    pdf.cell(right_width, 5, right_text, align='R')
    
    # 2. Reset X to left margin and Write Left Text
    pdf.set_xy(pdf.l_margin, y)
    pdf.set_font("Helvetica", style, 11)
    # We give it the full width minus the right text to avoid overlap if possible, or just full width
    pdf.cell(0, 5, left_text, align='L', ln=1) # ln=1 moves to next line

def markdown_to_pdf(markdown_text):
    """
    Converts simple Markdown text to a PDF file stream with improved formatting and link support.
    Parses resume structure for better layout (Centered Header, Right-aligned dates).
    """
    # Clean the input text first
    markdown_text = clean_markdown(markdown_text)
    
    pdf = PDF()
    pdf.add_page()
    # Set Narrow Margins (0.5 inch = 12.7 mm)
    pdf.set_margins(12.7, 12.7, 12.7)
    pdf.set_auto_page_break(auto=True, margin=12.7)
    
    # Use standard fonts
    pdf.set_font("Helvetica", size=10)
    
    lines = markdown_text.split('\n')
    
    # --- Resume Parser State ---
    is_header = True
    
    for line in lines:
        line = line.strip()
        
        # Skip empty lines
        if not line:
            if not is_header:
                pdf.ln(2) # Small space between items
            continue
            
        # Filter out unwanted text from LLM output
        if "Match Score:" in line or "Optimization Required" in line or "ATS-friendly resume" in line:
            continue
            
        # Sanitize line first
        line = sanitize_for_latin1(line)
        
        # Skip horizontal rules
        if line == '---' or line == '...':
            continue
            
        if line.startswith('## '):
            # Heading 2 - Starts a new section
            is_header = False
            pdf.ln(4)
            pdf.set_font("Helvetica", 'B', 12) # Section Header Size
            clean_content = line[2:].replace('**', '').upper()
            
            # Draw Section Title
            pdf.cell(0, 6, clean_content, new_x="LMARGIN", new_y="NEXT", align='L')
            
            # Draw a line under the section title
            y = pdf.get_y()
            pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
            pdf.ln(2)
            
            pdf.set_font("Helvetica", size=10) # Body font
            
        elif line.startswith('# '):
            # Heading 1 - Name (usually)
            pdf.set_font("Helvetica", 'B', 24) # Huge Name
            clean_content = line[2:].replace('**', '')
            # Center the Name
            pdf.cell(0, 10, clean_content, new_x="LMARGIN", new_y="NEXT", align='C')
            pdf.set_font("Helvetica", size=10) # Reset
            
        elif line.startswith('- ') or line.startswith('* '):
            # Bullet point
            clean_line = line[2:].replace('**', '')
            
            # Bullet logic
            current_y = pdf.get_y()
            
            # Draw bullet
            pdf.set_x(15) # Indent for bullet
            pdf.cell(5, 5, chr(149), align='R') 
            
            # Draw text
            pdf.set_xy(20, current_y)
            
            # Check for links in bullet points
            if '[' in clean_line and '](' in clean_line:
                 write_text_with_links(pdf, clean_line)
            else:
                 pdf.multi_cell(0, 5, clean_line)
                  
        else:
            # Normal text or Split Line (Company | Date)
            clean_line = line.replace('**', '').replace('*', '') # Remove markdown markers for clean text
            
            if is_header:
                # Center contact info in header
                # If the line contains pipes |, it's likely contact info "Email | Phone | LinkedIn"
                if '|' in clean_line:
                     # Replace | with spaced | for better look or just print
                     clean_line = clean_line.replace('|', '  |  ')
                
                pdf.cell(0, 5, clean_line, new_x="LMARGIN", new_y="NEXT", align='C')
            else:
                # Check if this is a "Split Line" candidate
                # Pattern 1: Explicit "|" separator
                if '|' in line:
                    parts = line.split('|')
                    if len(parts) >= 2:
                        # Take the last part as the "Right" side (Date/Location)
                        right = parts[-1].strip().replace('**', '')
                        # Join the rest as the "Left" side
                        left = " | ".join(parts[:-1]).strip().replace('**', '')
                        
                        # Check if original had bold on left
                        is_bold = '**' in parts[0]
                        draw_split_line(pdf, left, right, is_bold_left=is_bold)
                        continue
                        
                # Pattern 2: Ends with a date-like string (Simple heuristic: 4 digits at end)
                # Regex for date at end of line: " ... Jan 2020" or " ... 2020 - 2021"
                date_match = re.search(r'\s+((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|Summer|Winter|Spring|Fall|Present|Current|\d{4}).*?)$', clean_line)
                
                if date_match and len(clean_line) < 100: # Only apply to short lines (headers)
                    date_text = date_match.group(1)
                    left_text = clean_line[:date_match.start()].strip()
                    
                    # Determine styling
                    is_bold = '**' in line and line.startswith('**')
                    is_italic = '*' in line and line.startswith('*') and not line.startswith('**')
                    
                    draw_split_line(pdf, left_text, date_text, is_bold_left=is_bold, is_italic_left=is_italic)
                else:
                    # Just normal text
                    if '[' in clean_line and '](' in clean_line:
                        write_text_with_links(pdf, clean_line)
                    else:
                        # If it looks like a sub-header (bold), make it bold
                        if line.startswith('**') and line.endswith('**'):
                            pdf.set_font("Helvetica", 'B', 10)
                            pdf.multi_cell(0, 5, clean_line)
                            pdf.set_font("Helvetica", '', 10)
                        else:
                            pdf.multi_cell(0, 5, clean_line)
            
    pdf_stream = io.BytesIO()
    pdf.output(pdf_stream)
    pdf_stream.seek(0)
    return pdf_stream
